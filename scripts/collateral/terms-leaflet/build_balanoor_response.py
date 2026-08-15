from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT / "deliverables" / "balanoor-response"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "GardenSuite_Response_to_Balanoor.docx"

BLUE = RGBColor(31, 78, 120)
ACCENT = RGBColor(46, 117, 182)
MUTED = RGBColor(91, 104, 117)
INK = RGBColor(23, 33, 43)
LIGHT_BLUE = "EAF3FA"
LIGHT_ORANGE = "FCE4D6"
ORANGE = RGBColor(131, 60, 0)


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_label_detail(doc, label, detail):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    label_run = p.add_run(label + ": ")
    set_run_font(label_run, bold=True, color=BLUE)
    detail_run = p.add_run(detail)
    set_run_font(detail_run)
    return p


def add_callout(doc, title, text, fill=LIGHT_BLUE, title_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.35
    return table


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.different_first_page_header_footer = False
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.right_margin = Inches(1.0)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

doc.core_properties.title = "GardenSuite response to Balanoor Plantations and Industries Ltd."
doc.core_properties.subject = "Face attendance and green leaf weighing"
doc.core_properties.author = "Sarbani Associates"
doc.core_properties.keywords = "GardenSuite, Balanoor Plantations, attendance, green leaf weighing"

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

h1 = doc.styles["Heading 1"]
h1.font.name = "Calibri"
h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = ACCENT
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(5)
h1.paragraph_format.keep_with_next = True

list_style = doc.styles["List Bullet"]
list_style.font.name = "Calibri"
list_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
list_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
list_style.font.size = Pt(11)
list_style.font.color.rgb = INK

footer = section.footer
fp = footer.paragraphs[0]
fp.paragraph_format.space_before = Pt(0)
add_page_number(fp)

kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(2)
kr = kicker.add_run("TECHNICAL AND BUDGETARY RESPONSE")
set_run_font(kr, size=9, color=ACCENT, bold=True)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
tr = title.add_run("Face Attendance and Green Leaf Weighing")
set_run_font(tr, size=24, color=BLUE, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
sr = subtitle.add_run("Prepared for Balanoor Plantations and Industries Ltd.")
set_run_font(sr, size=12.5, color=MUTED)

add_label_detail(doc, "Subject", "GardenSuite quotation for face attendance and green leaf weighing")

add_body(doc, "Dear Sir/Madam,")
add_body(doc, "Thank you for contacting GardenSuite regarding the green leaf weighing and field attendance requirements of Balanoor Plantations and Industries Ltd.")
add_body(doc, "Please find attached our budgetary quotation, technical response and product brochure.")

add_callout(
    doc,
    "Proposed solution",
    "Our current solution combines a 100 kg Bluetooth smart hanging scale with the GardenSuite face attendance Android app. We currently offer one standard scale configuration. The number of scales depends on the number of kamjari or weighing points, while the software plan depends on the number of Android field devices."
)

add_body(doc, "Our response to your questions is given below:")

add_heading(doc, "Available models and pricing")
add_label_detail(doc, "GardenSuite Wireless Smart Hanging Scale, 100 kg", "Rs. 8,000 per scale, one-time")
add_label_detail(doc, "Face Attendance and Smart Weighing software", "Rs. 20,000 onwards per year, based on the number of Android devices")
add_label_detail(doc, "On-site installation and training", "Rs. 1,000 per day")
add_label_detail(doc, "Travel, accommodation and local transport", "Extra at actual cost. For travel beyond 500 km, return flight travel must be provided or reimbursed.")
add_body(doc, "All software plans include the face attendance app, Bluetooth scale connection, dashboard, support and software updates.")

add_heading(doc, "Technical specifications")
add_body(doc, "The standard offer is a 100 kg digital hanging scale with Bluetooth connection to the GardenSuite Android app. The scale reading is captured against the face-verified worker. The system can record gross weight, deductions and net weight. Manual weight entry is also available if the scale is temporarily unavailable.")
add_body(doc, "The final scale make, accuracy or graduation, battery specification and calibration documentation will be confirmed against the selected supply batch before the purchase order.")

add_heading(doc, "Fingerprint or facial attendance")
add_body(doc, "The current GardenSuite field solution supports facial recognition. Fingerprint or thumb attendance is not included.")

add_heading(doc, "How face attendance works")
add_body(doc, "Workers are first enrolled using face images linked to their employee codes. During field work, the supervisor uses an Android phone camera to verify the worker. The app records the attendance and local time.")
add_body(doc, "During plucking, the Bluetooth scale sends the leaf weight to the same verified worker record. This helps stop proxy attendance and reduces mistakes caused by paper weight chits.")

add_heading(doc, "Existing software changes and data exchange")
add_body(doc, "No other software is required for a standalone GardenSuite deployment. Existing GardenSuite installations can be configured with the Face Attendance and Smart Weighing module.")
add_callout(
    doc,
    "CSV export for your existing software",
    "GardenSuite can generate attendance and green leaf weighing data in CSV format. The CSV file can be imported into your existing payroll or ERP software, subject to that software's supported import format and capabilities. Any data mapping, import setup or integration within the existing software must be carried out by your current software vendor. We can provide a sample CSV file and explain the available fields to assist them.",
    fill=LIGHT_ORANGE,
    title_color=ORANGE,
)
add_body(doc, "If a different file structure or any custom output is required, we will first need to review the existing software version, required import format and sample attendance and weight files. Any GardenSuite-side customisation will be quoted separately after this review.")

add_heading(doc, "Network and internet requirement")
add_body(doc, "Attendance and weighing work offline in the field. Bluetooth communication between the scale and Android phone does not require internet.")
add_body(doc, "Internet is required later for data synchronisation, dashboard access, software updates and remote support.")

add_heading(doc, "Maintenance, support and warranty")
add_body(doc, "Phone, WhatsApp and remote desktop support are included with the active software subscription. Software updates are also included.")
add_body(doc, "The scale carries a one-year warranty for manufacturing defects. The final warranty procedure, exclusions and replacement terms will be stated in the final quotation and invoice.")
add_body(doc, "Initial setup and training can normally be planned for approximately three working days. The actual duration will depend on the number of weighing points, supervisors and any software integration work.")

add_heading(doc, "Starting and closing working times")
add_body(doc, "Yes. GardenSuite can record face-verified clock-in and clock-out times for general garden work. Work sessions also record their starting and closing times. Individual plucking and weighing records retain their capture time.")

add_heading(doc, "Additional requirements")
add_body(doc, "Balanoor Plantations will need to provide:")
for item in [
    "Compatible Android phones with front cameras and Bluetooth",
    "Charging and electricity arrangements",
    "Internet access at an office or regular synchronisation point",
    "Hanging frame or stand for the scale, unless separately included",
    "SIM or data connection, router, UPS or printer if required",
]:
    add_bullet(doc, item)

add_callout(
    doc,
    "Important: South India on-site service",
    "Sarbani Associates does not currently maintain a local on-site service team in South India. Regular GardenSuite service, training and support are provided online and off-site. The GardenSuite field application can continue working offline, with internet required later for synchronisation.",
    fill=LIGHT_ORANGE,
    title_color=ORANGE,
)
add_body(doc, "For installation, training, commissioning or any later on-site visit, our team will need to travel to the estate. The on-site service charge is Rs. 1,000 per day. In addition, actual return airfare, airport and local transport, suitable accommodation, boarding and meals for the visiting Sarbani Associates team must be arranged or reimbursed by Balanoor Plantations. These expenses are separate from the software, subscription, implementation and service charges.")

add_heading(doc, "Information required for the formal quotation")
add_body(doc, "To prepare the final quantity-based quotation, please confirm:")
for item in [
    "Number of estates or divisions",
    "Number of kamjari or weighing points",
    "Number of supervisors and Android phones",
    "Approximate workers handled per day",
    "Existing payroll or ERP software details",
    "Required clock-in, clock-out and reporting rules",
    "Estate location and nearest airport",
    "Two suitable dates and times for an online demonstration",
]:
    add_bullet(doc, item)

add_callout(
    doc,
    "Next step",
    "Once we receive the above information, we can prepare and send a formal quotation with the final quantities, scope, pricing and applicable terms."
)

add_body(doc, "We look forward to demonstrating the complete workflow to your team.")
add_body(doc, "Regards,")

signature = doc.add_paragraph()
signature.paragraph_format.space_after = Pt(0)
for text, size, bold, color in [
    ("Kaushik Majumder", 11.5, True, BLUE),
    ("GardenSuite by Sarbani Associates", 10.5, True, ACCENT),
    ("Bagdogra, Siliguri, West Bengal", 10, False, MUTED),
    ("Phone: 9734101330", 10, False, MUTED),
    ("Email: sarbaniassociates@gmail.com", 10, False, MUTED),
    ("Website: https://gardensuite.in", 10, False, MUTED),
]:
    r = signature.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    r.add_break()

doc.save(OUT_PATH)
print(OUT_PATH)
